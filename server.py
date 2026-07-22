#!/usr/bin/env python3
import datetime as dt
from email.parser import BytesParser
from email.policy import default
import json
import os
import re
import sqlite3
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from xml.etree import ElementTree as ET
from zipfile import ZipFile


ROOT = Path(__file__).resolve().parent
DB_PATH = Path(os.environ.get("SOW_BUILDER_DB", ROOT / "sow_builder.db"))
TEMPLATE_DOCX = Path("/Users/adrianchatto/Downloads/Informa-SOW-draft (4).docx")
COVER_LOGO = ROOT / "gstt-template-assets" / "image1.png"
BRAND_PURPLE = "5B2EE6"
BRAND_BLUE = "0877EA"
BRAND_PINK = "EF0F64"
BRAND_ORANGE = "F57C00"
BRAND_GREEN = "0AA85A"
BRAND_INK = "17202A"
BRAND_LINE = "DFE4EA"
PHASE_COLOURS = [BRAND_PURPLE, BRAND_PINK, BRAND_ORANGE, BRAND_GREEN, "6B35D4", BRAND_BLUE, "0D7EEA"]


def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS sows (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                opportunity_number TEXT,
                customer TEXT,
                project TEXT,
                payload TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        conn.commit()


def json_response(handler, payload, status=200):
    body = json.dumps(payload).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def parse_uploaded_file(handler):
    content_type = handler.headers.get("Content-Type", "")
    length = int(handler.headers.get("Content-Length", "0"))
    body = handler.rfile.read(length)
    raw = (
        f"Content-Type: {content_type}\r\n"
        "MIME-Version: 1.0\r\n\r\n"
    ).encode("utf-8") + body
    message = BytesParser(policy=default).parsebytes(raw)
    if not message.is_multipart():
        raise ValueError("Expected multipart upload")
    for part in message.iter_parts():
        if part.get_content_disposition() != "form-data":
            continue
        if part.get_param("name", header="content-disposition") != "file":
            continue
        filename = part.get_filename()
        if not filename:
            break
        return filename, part.get_payload(decode=True) or b""
    raise ValueError("Missing file")


def extract_pptx(data):
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    output = []
    with ZipFile(BytesIO(data)) as deck:
        slides = sorted(
            [name for name in deck.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")],
            key=lambda name: int(name.rsplit("slide", 1)[1].split(".xml")[0]),
        )
        for index, slide in enumerate(slides, 1):
            root = ET.fromstring(deck.read(slide))
            texts = [node.text.strip() for node in root.findall(".//a:t", ns) if node.text and node.text.strip()]
            if texts:
                output.append(f"--- SLIDE {index} ---\n" + "\n".join(texts))
    return "\n\n".join(output)


def extract_pdf(data):
    try:
        from pypdf import PdfReader
    except Exception:
        return "PDF uploaded. Install pypdf in this Python environment for PDF text extraction."

    reader = PdfReader(BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- PAGE {index} ---\n{text.strip()}")
    return "\n\n".join(pages)


def extract_text(filename, data):
    suffix = Path(filename).suffix.lower()
    if suffix == ".pptx":
        return extract_pptx(data)
    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    return f"{filename} uploaded. This prototype extracts .pptx, .pdf, .txt and .md files."


def safe_name(value):
    return re.sub(r"[^A-Za-z0-9]+", "-", value or "Customer").strip("-") or "Customer"


def optional_enabled(data, key):
    return data.get("optionalSections", {}).get(key) == "include"


def document_name(data):
    customer = (data.get("customer") or "").strip()
    project = (data.get("project") or "").strip()
    prefix = "" if project.lower().startswith(customer.lower()) else f"{customer} - "
    return f"{prefix}{project} - CloudInteract SOW"


def cover_project_name(data):
    customer = (data.get("customer") or "").strip()
    project = (data.get("project") or "").strip()
    return project if project.lower().startswith(customer.lower()) else f"{customer} - {project}"


def plan_week_count(data):
    try:
        return max(1, min(52, int(data.get("duration") or 10)))
    except (TypeError, ValueError):
        return 10


def render_poap_jpeg(phases, week_count=10):
    from PIL import Image, ImageDraw, ImageFont

    week_count = max(1, min(52, int(week_count or 10)))
    width, height = 1800, 1080
    left, top = 24, 28
    label_width = 260
    grid_width = 1770
    week_width = (grid_width - label_width) / week_count
    header_height, row_height = 78, 124
    grid_height = header_height + row_height * len(phases)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    def font(size, bold=False):
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                pass
        return ImageFont.load_default()

    f_header = font(24, True)
    f_phase = font(24, True)
    f_small_bold = font(18, True)
    f_small = font(18)
    f_bar = font(22, True)
    f_mile = font(18, True)
    f_mile_small = font(17)

    def text_center(x, y, text, fill, fnt):
        bbox = draw.textbbox((0, 0), text, font=fnt)
        draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2), text, fill=fill, font=fnt)

    def wrap_text(text, x, y, max_width, line_height, fnt, fill, max_lines=3, center=False):
        words = str(text).split()
        lines, line = [], ""
        for word in words:
            candidate = f"{line} {word}".strip()
            bbox = draw.textbbox((0, 0), candidate, font=fnt)
            if bbox[2] - bbox[0] > max_width and line:
                lines.append(line)
                line = word
            else:
                line = candidate
        if line:
            lines.append(line)
        for index, item in enumerate(lines[:max_lines]):
            tx = x
            if center:
                bbox = draw.textbbox((0, 0), item, font=fnt)
                tx = x + (max_width - (bbox[2] - bbox[0])) / 2
            draw.text((tx, y + index * line_height), item, fill=fill, font=fnt)

    draw.rounded_rectangle((left, top, left + grid_width, top + grid_height), radius=18, fill="white", outline="#dfe4ea", width=2)
    text_center(left + label_width / 2, top + 40, "PHASE", "#111827", f_header)
    for i in range(week_count):
        text_center(left + label_width + i * week_width + week_width / 2, top + 40, f"WEEK {i + 1}", "#111827", f_header)
    for i in range(1, week_count + 1):
        x = left + label_width + i * week_width
        draw.line((x, top, x, top + grid_height), fill="#dfe4ea", width=2)
    for r in range(len(phases) + 1):
        y = top + header_height + r * row_height
        draw.line((left, y, left + grid_width, y), fill="#dfe4ea", width=2)

    symbols = ["?", "✎", "⚙", "▥", "✓", "★", "◆"]
    for index, phase in enumerate(phases):
        colour = "#" + PHASE_COLOURS[index % len(PHASE_COLOURS)]
        y = top + header_height + index * row_height
        draw.ellipse((left + 20, y + 30, left + 84, y + 94), fill=colour)
        text_center(left + 52, y + 63, symbols[index % len(symbols)], "white", f_bar)
        wrap_text(phase.get("name", ""), left + 96, y + 26, 132, 25, f_phase, "#111827", 2)
        draw.text((left + 96, y + 72), phase.get("sprint", ""), fill=colour, font=f_small_bold)
        draw.text((left + 96, y + 96), f"{phase.get('weeks', '')} Weeks", fill="#111827", font=f_small)

        start = int(phase.get("start", 1))
        duration = int(phase.get("weeks", 1))
        span = min(duration, week_count + 1 - start)
        if span < 1 or start > week_count:
            continue
        bar_x = left + label_width + (start - 1) * week_width + 8
        bar_y = y + 42
        bar_w = span * week_width - 16
        draw.rounded_rectangle((bar_x, bar_y, bar_x + bar_w, bar_y + 52), radius=13, fill=colour)
        text_center(bar_x + bar_w / 2, bar_y + 27, f"{duration} {'week' if duration == 1 else 'weeks'}", "white", f_bar)

    milestone_top = top + grid_height + 58
    milestone_gap = grid_width / max(len(phases), 1)
    for index, phase in enumerate(phases):
        colour = "#" + PHASE_COLOURS[index % len(PHASE_COLOURS)]
        x = left + index * milestone_gap + milestone_gap / 2
        draw.ellipse((x - 31, milestone_top - 31, x + 31, milestone_top + 31), fill=colour)
        text_center(x, milestone_top + 1, symbols[index % len(symbols)], "white", f_bar)
        if index < len(phases) - 1:
            text_center(x + milestone_gap / 2, milestone_top + 2, "›", "#8f98a5", font(42))
        wrap_text("Kick-off" if index == 0 else phase.get("name", ""), x - 70, milestone_top + 48, 140, 20, f_mile, "#111827", 2, center=True)
        wrap_text(phase.get("milestone", ""), x - 74, milestone_top + 94, 148, 21, f_mile_small, "#4b5563", 3, center=True)

    buffer = BytesIO()
    image.save(buffer, format="JPEG", quality=92)
    return buffer.getvalue()


def render_cover_jpeg(data):
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1284, 1800
    image = Image.new("RGB", (width, height), "#F20B6F")
    draw = ImageDraw.Draw(image, "RGBA")

    def font(size, bold=False, serif=False):
        candidates = [
            "/System/Library/Fonts/Supplemental/Times New Roman.ttf" if serif else None,
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
        ]
        for candidate in [item for item in candidates if item]:
            try:
                return ImageFont.truetype(candidate, size)
            except Exception:
                pass
        return ImageFont.load_default()

    for y in range(height):
        t = y / max(height - 1, 1)
        if t < 0.45:
            local = t / 0.45
            start, end = (255, 139, 61), (251, 92, 88)
        else:
            local = (t - 0.45) / 0.55
            start, end = (251, 92, 88), (242, 11, 111)
        colour = tuple(int(start[i] + (end[i] - start[i]) * local) for i in range(3))
        draw.line((0, y, width, y), fill=colour)

    if COVER_LOGO.exists():
        logo = Image.open(COVER_LOGO).convert("RGBA")
        logo.thumbnail((430, 265), Image.LANCZOS)
        image.paste(logo, (116, 140), logo)
    else:
        draw.text((300, 210), "Cloud", fill="white", font=font(74, True))
        draw.text((306, 285), "Interact", fill="white", font=font(36))

    def wrap(text, x, y, max_width, line_height, fnt, fill, max_lines=5):
        words = str(text).split()
        lines, line = [], ""
        for word in words:
            candidate = f"{line} {word}".strip()
            bbox = draw.textbbox((0, 0), candidate, font=fnt)
            if bbox[2] - bbox[0] > max_width and line:
                lines.append(line)
                line = word
            else:
                line = candidate
        if line:
            lines.append(line)
        for index, item in enumerate(lines[:max_lines]):
            draw.text((x, y + index * line_height), item, fill=fill, font=fnt)

    wrap(cover_project_name(data), 28, 960, 900, 118, font(92), "white", max_lines=4)
    draw.text((34, 1355), data.get("subtitle") or "Statement of Work", fill="white", font=font(34))

    buffer = BytesIO()
    image.save(buffer, format="JPEG", quality=94)
    return buffer.getvalue()


def version_rows(data):
    return [
        ["Version", "Date Modified", "Description", "Modified By"],
        [
            data.get("version") or "v0.1",
            "21 Jul 2026",
            data.get("versionComment") or "First draft",
            data.get("author") or data.get("supplier") or "CloudInteract",
        ],
    ]


def document_detail_rows(data):
    return [
        ["Client", data.get("customer") or ""],
        ["Date Created", "21 Jul 2026"],
        ["Document Type", "Statement of Work (SOW)"],
        ["Document Name", data.get("project") or ""],
        ["Author", data.get("author") or data.get("supplier") or "CloudInteract"],
    ]


def split_lines(value):
    return [line.strip() for line in str(value or "").splitlines() if line.strip()]


def dynamic_text(value, data):
    return (
        str(value or "")
        .replace("{customer}", data.get("customer") or "Client")
        .replace("{supplier}", data.get("supplier") or "CloudInteract")
        .replace("Informa", data.get("customer") or "Client")
    )


def document_file_base(data):
    parts = [data.get("opportunityNumber"), data.get("customer"), data.get("project")]
    return "-".join([safe_name(part) for part in parts if str(part or "").strip()])


def toc_rows(data):
    rows = [
        ["1", "Agreement"],
        ["2", "Background"],
        ["3", "Scope"],
        ["4", "Project Plan"],
        ["5", "Deliverables And Acceptance Criteria"],
        ["6", "Success Metrics"],
        ["7", "Dependencies"],
        ["8", "Change Control"],
        ["9", "Data Protection, Security And AI Governance"],
    ]
    optional = [
        ("ipDetail", "10", "Intellectual Property And Reuse"),
        ("confidentiality", "11", "Confidentiality"),
        ("assurance", "12", "Assurance And Oversight"),
        ("stage2", "13", "Stage 2 Roadmap"),
        ("legalBoilerplate", "14", "Extended Legal Boilerplate"),
    ]
    for key, number, title in optional:
        if optional_enabled(data, key):
            rows.append([number, title])
    if raci_rows(data):
        rows.append(["", "RACI"])
    rows.extend([["", "Commercials"], ["", "Signature"]])
    return rows


def commercial_rows(data):
    rows = []
    for row in data.get("commercialMilestones") or []:
        title = str(row.get("title") or "").strip()
        amount = str(row.get("amount") or "").strip()
        if title or amount:
            rows.append([title or "TBD", amount or "TBD"])
    if not rows:
        rows.append(["Fixed price", str(data.get("price") or "").strip() or "TBD"])
    discount = str(data.get("commercialDiscount") or "").strip()
    subtotal = sum_commercial_amounts(rows)
    if subtotal is not None and (len(rows) > 1 or discount):
        rows.append(["Subtotal", format_commercial_amount(subtotal)])
    if discount:
        discount_amount = parse_commercial_amount(discount)
        rows.append(["Discount", discount if discount_amount is None else f"-{format_commercial_amount(discount_amount)}"])
    if subtotal is not None:
        discount_amount = parse_commercial_amount(discount) if discount else 0
        rows.append(["Total", format_commercial_amount(max(0, subtotal - (discount_amount or 0)))])
    return [["Milestone", "Amount"]] + rows


def raci_rows(data):
    rows = []
    for row in data.get("raciRows") or []:
        values = [
            str(row.get("activity") or "").strip(),
            str(row.get("dateRange") or "").strip(),
            str(row.get("responsible") or "").strip(),
            str(row.get("accountable") or "").strip(),
            str(row.get("consulted") or "").strip(),
            str(row.get("informed") or "").strip(),
        ]
        if any(values):
            rows.append([value or "TBD" for value in values])
    if not rows:
        return []
    return [["Activity / deliverable", "Date range", "Responsible", "Accountable", "Consulted", "Informed"]] + rows


def parse_commercial_amount(value):
    raw = str(value or "").strip()
    if not raw or re.search(r"\btbd\b", raw, re.I):
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", raw.replace(",", ""))
    return abs(float(match.group(0))) if match else None


def sum_commercial_amounts(rows):
    total = 0.0
    numeric_rows = 0
    for _, amount in rows:
        value = parse_commercial_amount(amount)
        if value is None:
            continue
        total += value
        numeric_rows += 1
    return total if numeric_rows else None


def format_commercial_amount(value):
    amount = float(value)
    if amount.is_integer():
        return f"£{int(amount):,}"
    return f"£{amount:,.2f}"


def add_word_toc(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    title = doc.add_paragraph()
    title_run = title.add_run("Table of Contents")
    title_run.bold = True
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor(23, 32, 42)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run("Right-click and update field to refresh the table of contents.")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_docx(data):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    if "Normal" in [style.name for style in styles]:
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)]:
        if style_name not in [style.name for style in styles]:
            continue
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = RGBColor(23, 32, 42)
        style.font.size = Pt(size)

    def shade(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def remove_table_borders(table):
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "nil")

    def format_table(table):
        remove_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

    def set_cell_text(cell, value, bold=False):
        cell.text = str(value)
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def add_table(rows, bordered=False):
        table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        if bordered:
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
        else:
            format_table(table)
        for r, row in enumerate(rows):
            for c, cell_value in enumerate(row):
                set_cell_text(table.cell(r, c), cell_value, r == 0)
        return table

    def add_paragraphs(text):
        for line in split_lines(dynamic_text(text, data)):
            doc.add_paragraph(line)

    def add_bullets(text):
        for line in split_lines(dynamic_text(text, data)):
            doc.add_paragraph(line, style="List Bullet")

    def dt(text):
        return dynamic_text(text, data)

    def add_front_title(text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(17)
        run.font.color.rgb = RGBColor(23, 32, 42)
        return paragraph

    doc.add_picture(BytesIO(render_cover_jpeg(data)), width=Inches(7.05))

    doc.add_page_break()
    add_front_title("Document Details")
    add_table(document_detail_rows(data))
    add_front_title("Document Revision History")
    add_table(version_rows(data))
    for _ in range(14):
        doc.add_paragraph("")
    add_front_title("Proprietary Notice")
    add_paragraphs(data.get("proprietaryNotice"))
    doc.add_page_break()

    add_word_toc(doc)
    doc.add_page_break()

    doc.add_heading("1 Agreement", level=1)
    add_paragraphs(data.get("agreementText"))
    doc.add_heading("2 Background", level=1)
    doc.add_heading("2.1 Customer Overview", level=2)
    add_paragraphs(data.get("backgroundOverview"))
    doc.add_heading("2.2 Customer Requirements", level=2)
    add_paragraphs(data.get("backgroundRequirements"))
    doc.add_heading("3 Scope", level=1)
    doc.add_heading("3.1 Included", level=2)
    add_bullets(data.get("scopeIncluded"))
    doc.add_heading("3.2 Boundaries And Exclusions", level=2)
    add_bullets(data.get("scopeExclusions"))
    doc.add_heading("4 Project Plan", level=1)
    doc.add_heading("4.1 Plan On A Page", level=2)
    doc.add_paragraph("The plan below is indicative and assumes timely access, data, stakeholder availability and governance decisions.")
    doc.add_picture(BytesIO(render_poap_jpeg(data.get("phases", []), plan_week_count(data))), width=Inches(7.0))
    doc.add_heading("5 Deliverables And Acceptance Criteria", level=1)
    add_table([
        ["Deliverable", "Acceptance basis"],
        [dt("Working reasoning-agent prototype in {customer} AWS sandbox"), "Demonstrated against the selected software-request journey using dev ServiceNow and representative knowledge sources."],
        ["Microsoft Teams colleague experience", "Stakeholders can request software in plain language and receive resolution, approval routing or escalation updates in Teams."],
        ["ServiceNow and knowledge integrations", "Prototype can check entitlement, duplicate requests and route/document work through supported public APIs."],
        ["Handover pack and Stage 2 recommendations", dt("{customer} IT receives architecture notes, run considerations, backlog and recommended pilot/production next steps.")],
    ], bordered=True)
    doc.add_heading("6 Success Metrics", level=1)
    add_bullets(data.get("successMeasuresText"))
    doc.add_heading("7 Dependencies", level=1)
    add_bullets(data.get("dependenciesText"))
    doc.add_heading("8 Change Control", level=1)
    add_paragraphs(data.get("changeControlText"))
    doc.add_heading("9 Data Protection, Security And AI Governance", level=1)
    add_paragraphs(data.get("securityText"))
    rows = raci_rows(data)
    if rows:
        doc.add_heading("RACI", level=1)
        add_table(rows, bordered=True)
    doc.add_heading("Commercials", level=1)
    add_table(commercial_rows(data), bordered=True)
    doc.add_heading("Signature", level=1)
    add_table([[f"For {data.get('customer', '')}", f"For {data.get('supplier', '')}"], ["Name:\n\nTitle:\n\nDate:", "Name:\n\nTitle:\n\nDate:"]], bordered=True)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def add_docx_plan(doc, phases, shade, week_count=10):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches

    week_count = max(1, min(52, int(week_count or 10)))
    weeks = list(range(1, week_count + 1))
    table = doc.add_table(rows=len(phases) + 1, cols=week_count + 1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.cell(0, 0).text = "Phase"
    shade(table.cell(0, 0), "F6F8FB")
    for i, week in enumerate(weeks, 1):
        table.cell(0, i).text = f"W{week}"
        shade(table.cell(0, i), "F6F8FB")
    for r, phase in enumerate(phases, 1):
        table.cell(r, 0).text = f"{phase.get('name', '')}\n{phase.get('sprint', '')}\n{phase.get('weeks', '')} weeks"
        colour = PHASE_COLOURS[(r - 1) % len(PHASE_COLOURS)]
        start = int(phase.get("start", 1))
        duration = int(phase.get("weeks", 1))
        for week in weeks:
            cell = table.cell(r, week)
            if start <= week < start + duration:
                cell.text = f"{duration} week" + ("" if duration == 1 else "s") if week == start else ""
                shade(cell, colour)
            else:
                cell.text = ""
    doc.add_paragraph("Milestones: " + " | ".join([f"{p.get('name', '')}: {p.get('milestone', '')}" for p in phases]))


def build_pdf(data):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, PageBreak, SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=14 * mm, leftMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Brand", parent=styles["Normal"], textColor=colors.HexColor("#5B2EE6"), fontName="Helvetica-Bold", fontSize=14, spaceAfter=14))
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=8, leading=10))
    story = [Image(BytesIO(render_cover_jpeg(data)), width=182 * mm, height=255 * mm), PageBreak()]

    def heading(text, level=2):
        story.append(Spacer(1, 8))
        story.append(Paragraph(escape(text), styles["Heading2" if level == 2 else "Heading3"]))

    def paras(text):
        for line in split_lines(dynamic_text(text, data)):
            story.append(Paragraph(escape(line), styles["BodyText"]))

    def bullets(text):
        for line in split_lines(dynamic_text(text, data)):
            story.append(Paragraph(escape(line), styles["Bullet"]))

    def dt(text):
        return dynamic_text(text, data)

    heading("Document Details")
    story.append(make_pdf_table(document_detail_rows(data), None))
    heading("Document Revision History")
    story.append(make_pdf_table(version_rows(data), None))
    story.append(Spacer(1, 92 * mm))
    heading("Proprietary Notice")
    paras(data.get("proprietaryNotice"))
    story.append(PageBreak())

    heading("Table of Contents")
    story.append(make_pdf_table(toc_rows(data), [18 * mm, 152 * mm]))
    story.append(PageBreak())

    heading("1 Agreement")
    paras(data.get("agreementText"))
    heading("2 Background")
    heading("2.1 Customer Overview", level=3)
    paras(data.get("backgroundOverview"))
    heading("2.2 Customer Requirements", level=3)
    paras(data.get("backgroundRequirements"))
    heading("3 Scope")
    heading("3.1 Included", level=3)
    bullets(data.get("scopeIncluded"))
    heading("3.2 Boundaries And Exclusions", level=3)
    bullets(data.get("scopeExclusions"))
    heading("4 Project Plan")
    heading("4.1 Plan On A Page", level=3)
    story.append(Paragraph("The plan below is indicative and assumes timely access, data, stakeholder availability and governance decisions.", styles["BodyText"]))
    story.append(Image(BytesIO(render_poap_jpeg(data.get("phases", []), plan_week_count(data))), width=180 * mm, height=108 * mm))
    heading("5 Deliverables And Acceptance Criteria")
    story.append(make_pdf_table([
        ["Deliverable", "Acceptance basis"],
        [dt("Working reasoning-agent prototype in {customer} AWS sandbox"), "Demonstrated against the selected software-request journey using dev ServiceNow and representative knowledge sources."],
        ["Microsoft Teams colleague experience", "Stakeholders can request software in plain language and receive resolution, approval routing or escalation updates in Teams."],
        ["ServiceNow and knowledge integrations", "Prototype can check entitlement, duplicate requests and route/document work through supported public APIs."],
        ["Handover pack and Stage 2 recommendations", dt("{customer} IT receives architecture notes, run considerations, backlog and recommended pilot/production next steps.")],
    ], None, bordered=True))
    heading("6 Success Metrics")
    bullets(data.get("successMeasuresText"))
    heading("7 Dependencies")
    bullets(data.get("dependenciesText"))
    heading("8 Change Control")
    paras(data.get("changeControlText"))
    heading("9 Data Protection, Security And AI Governance")
    paras(data.get("securityText"))

    optional_items = [
        ("ipDetail", "10 Intellectual Property And Reuse", "Unless otherwise agreed in signed commercial terms, the customer owns customer-specific outputs and data supplied by the customer. CloudInteract retains ownership of pre-existing tools, methods, know-how, templates and reusable accelerators used to deliver the work."),
        ("confidentiality", "11 Confidentiality", "Each party will protect confidential information disclosed in connection with this SOW. Where a separate NDA or master agreement exists, that agreement will take precedence over this summary wording."),
        ("assurance", "12 Assurance And Oversight", "The prototype will include outcome verification, human review for uncertain or failed cases, audit logging for agent decisions and visibility of agreed success metrics. Any standing auto-approval rule must be agreed before use."),
        ("stage2", "13 Stage 2 Roadmap", "Following Stage 1, the recommended path is a controlled pilot with a friendly cohort, then production hardening, phased rollout by segment, additional channels and expansion into further use cases."),
        ("legalBoilerplate", "14 Extended Legal Boilerplate", "Final legal terms should confirm warranty, liability, assignment, third-party rights, counterparts, governing law and jurisdiction, either in this SOW or in the governing master services agreement."),
    ]
    for key, title, text in optional_items:
        if optional_enabled(data, key):
            heading(title)
            story.append(Paragraph(escape(text), styles["BodyText"]))

    rows = raci_rows(data)
    if rows:
        heading("RACI")
        story.append(make_pdf_table(rows, None, bordered=True))

    heading("Commercials")
    story.append(make_pdf_table(commercial_rows(data), None, bordered=True))
    heading("Signature")
    story.append(make_pdf_table([[f"For {data.get('customer', '')}", f"For {data.get('supplier', '')}"], ["Name:\n\nTitle:\n\nDate:", "Name:\n\nTitle:\n\nDate:"]], None, bordered=True))

    doc.build(story, onFirstPage=lambda canvas, doc: None, onLaterPages=pdf_footer)
    return buffer.getvalue()


def improve_with_openrouter(text):
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        raise RuntimeError("OPENROUTER_API_KEY is not set on the server.")
    payload = {
        "model": os.environ.get("OPENROUTER_MODEL", "anthropic/claude-3.5-sonnet"),
        "messages": [
            {
                "role": "system",
                "content": "Improve this Statement of Work section. Keep the commercial/legal meaning, use concise professional UK English, and return only the improved text.",
            },
            {"role": "user", "content": text},
        ],
        "temperature": 0.2,
    }
    request = Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://127.0.0.1:5173",
            "X-Title": "SOW Builder",
        },
        method="POST",
    )
    with urlopen(request, timeout=45) as response:
        result = json.loads(response.read().decode("utf-8"))
    return result["choices"][0]["message"]["content"].strip()


def make_pdf_table(rows, widths, bordered=False):
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import Table, TableStyle

    if not widths:
        column_count = max(len(row) for row in rows)
        widths = [170 * mm / column_count] * column_count
    table = Table(rows, colWidths=widths, hAlign="LEFT")
    commands = [
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if bordered:
        commands.extend([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD3DF")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F6F8FB")),
        ])
    else:
        commands.append(("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#EEF1F5")))
    table.setStyle(TableStyle(commands))
    return table


def make_pdf_plan(phases, week_count=10):
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import Table, TableStyle

    week_count = max(1, min(52, int(week_count or 10)))
    rows = [["Phase"] + [f"W{i}" for i in range(1, week_count + 1)]]
    styles = [
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DFE4EA")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F6F8FB")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    for r, phase in enumerate(phases, 1):
        duration = int(phase.get("weeks", 1))
        start = int(phase.get("start", 1))
        rows.append([f"{phase.get('name', '')}\n{phase.get('sprint', '')}"] + [""] * week_count)
        colour = colors.HexColor("#" + PHASE_COLOURS[(r - 1) % len(PHASE_COLOURS)])
        for week in range(start, min(week_count + 1, start + duration)):
            rows[r][week] = f"{duration}w" if week == start else ""
            styles.append(("BACKGROUND", (week, r), (week, r), colour))
            styles.append(("TEXTCOLOR", (week, r), (week, r), colors.white))
    table = Table(rows, colWidths=[35 * mm] + [13.5 * mm] * 10, rowHeights=[8 * mm] + [12 * mm] * len(phases), hAlign="LEFT")
    table.setStyle(TableStyle(styles))
    return table


def pdf_footer(canvas, doc):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#5B2EE6"))
    canvas.setLineWidth(1)
    canvas.line(14 * mm, 282 * mm, 196 * mm, 282 * mm)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#687385"))
    canvas.drawString(14 * mm, 8 * mm, "CloudInteract Confidential")
    canvas.drawRightString(196 * mm, 8 * mm, f"Page {doc.page}")
    canvas.restoreState()


class Handler(SimpleHTTPRequestHandler):
    def translate_path(self, path):
        parsed = urlparse(path).path.lstrip("/")
        return str((ROOT / parsed).resolve())

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/api/sows":
            self.handle_sow_list()
            return
        if parsed.path.startswith("/api/sows/"):
            self.handle_sow_get(parsed.path.rsplit("/", 1)[-1])
            return
        super().do_GET()

    def do_DELETE(self):
        parsed = urlparse(self.path)
        if parsed.path.startswith("/api/sows/"):
            self.handle_sow_delete(parsed.path.rsplit("/", 1)[-1])
            return
        self.send_error(404)

    def do_POST(self):
        if self.path.startswith("/api/export/"):
            self.handle_export()
            return

        if self.path == "/api/sows":
            self.handle_sow_save()
            return

        if self.path == "/api/improve":
            self.handle_improve()
            return

        if self.path != "/api/analyse":
            self.send_error(404)
            return

        try:
            filename, data = parse_uploaded_file(self)
            text = extract_text(filename, data)
            payload = {"filename": filename, "text": text}
            self.send_response(200)
        except Exception as exc:
            payload = {"error": str(exc)}
            self.send_response(500)

        body = json.dumps(payload).encode("utf-8")
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def read_json_body(self):
        length = int(self.headers.get("Content-Length", "0"))
        return json.loads(self.rfile.read(length).decode("utf-8") or "{}")

    def handle_improve(self):
        try:
            payload = self.read_json_body()
            improved = improve_with_openrouter(payload.get("text", ""))
            json_response(self, {"text": improved})
        except Exception as exc:
            json_response(self, {"error": str(exc)}, 500)

    def handle_sow_list(self):
        init_db()
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute(
                """
                SELECT id, opportunity_number, customer, project, updated_at
                FROM sows
                ORDER BY updated_at DESC
                """
            ).fetchall()
        json_response(
            self,
            {
                "items": [
                    {
                        "id": row["id"],
                        "opportunityNumber": row["opportunity_number"],
                        "customer": row["customer"],
                        "project": row["project"],
                        "updatedAt": row["updated_at"],
                    }
                    for row in rows
                ]
            },
        )

    def handle_sow_get(self, item_id):
        init_db()
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM sows WHERE id = ?", (item_id,)).fetchone()
        if not row:
            json_response(self, {"error": "Saved SOW not found"}, 404)
            return
        json_response(
            self,
            {
                "item": {
                    "id": row["id"],
                    "opportunityNumber": row["opportunity_number"],
                    "customer": row["customer"],
                    "project": row["project"],
                    "updatedAt": row["updated_at"],
                    "data": json.loads(row["payload"]),
                }
            },
        )

    def handle_sow_save(self):
        init_db()
        try:
            data = self.read_json_body()
            now = dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()
            with sqlite3.connect(DB_PATH) as conn:
                conn.execute(
                    """
                    INSERT INTO sows (opportunity_number, customer, project, payload, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (
                        data.get("opportunityNumber"),
                        data.get("customer"),
                        data.get("project"),
                        json.dumps(data),
                        now,
                        now,
                    ),
                )
                item_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
                conn.commit()
            json_response(self, {"id": item_id, "updatedAt": now})
        except Exception as exc:
            json_response(self, {"error": str(exc)}, 500)

    def handle_sow_delete(self, item_id):
        init_db()
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("DELETE FROM sows WHERE id = ?", (item_id,))
            conn.commit()
        json_response(self, {"ok": True})

    def handle_export(self):
        format_name = self.path.rsplit("/", 1)[-1]
        try:
            data = self.read_json_body()
            if format_name == "docx":
                body = build_docx(data)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                extension = "docx"
            elif format_name == "pdf":
                body = build_pdf(data)
                content_type = "application/pdf"
                extension = "pdf"
            else:
                self.send_error(404)
                return
        except Exception as exc:
            payload = json.dumps({"error": str(exc)}).encode("utf-8")
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            self.wfile.write(payload)
            return

        filename = f"{document_file_base(data)}.{extension}"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


if __name__ == "__main__":
    init_db()
    host = os.environ.get("SOW_BUILDER_HOST", "127.0.0.1")
    port = int(os.environ.get("SOW_BUILDER_PORT", "5173"))
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"SOW Builder running at http://{host}:{port}")
    server.serve_forever()
